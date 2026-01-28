from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
from docx2pdf import convert






def create_circular_image(image_path, output_path=None, border_width=0, border_color=(255, 255, 255, 255)):
    """
    Создает идеально круглое изображение с возможностью добавления обводки
    
    Args:
        image_path (str): Путь к исходному изображению
        output_path (str, optional): Путь для сохранения результата. Если None, будет создан автоматически
        border_width (int): Ширина обводки в пикселях (0 - без обводки)
        border_color (tuple): Цвет обводки в формате RGBA
        
    Returns:
        str: Путь к сохраненному круглому изображению
    """
    try:
        # Открываем изображение
        img = Image.open(image_path).convert("RGBA")
        print(f"✓ Загружено изображение: {image_path} ({img.size[0]}x{img.size[1]})")
        
        # Создаем квадратное изображение
        width, height = img.size
        size = min(width, height)  # Размер квадрата
        
        # Координаты для обрезки до центрального квадрата
        left = (width - size) // 2
        top = (height - size) // 2
        right = left + size
        bottom = top + size
        
        # Обрезаем до квадрата
        img_square = img.crop((left, top, right, bottom))
        
        # Создаем маску для идеального круга
        mask = Image.new('L', (size, size), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse([0, 0, size, size], fill=255)
        
        # Применяем маску
        circular_img = Image.new('RGBA', (size, size), (0, 0, 0, 0))
        circular_img.paste(img_square, (0, 0), mask=mask)
        
        # Добавляем обводку (если нужно)
        if border_width > 0:
            total_size = size + border_width * 2
            bordered_img = Image.new('RGBA', (total_size, total_size), (0, 0, 0, 0))
            
            # Создаем маску для обводки
            border_mask = Image.new('L', (total_size, total_size), 0)
            draw_border = ImageDraw.Draw(border_mask)
            draw_border.ellipse([0, 0, total_size, total_size], fill=255)
            
            # Заливаем обводку цветом
            border_layer = Image.new('RGBA', (total_size, total_size), border_color)
            bordered_img.paste(border_layer, (0, 0), mask=border_mask)
            
            # Вставляем основное изображение
            bordered_img.paste(circular_img, (border_width, border_width), mask=circular_img)
            circular_img = bordered_img
            size = total_size
        
        # Сохраняем результат
        if output_path is None:
            name, ext = os.path.splitext(image_path)
            output_path = f"circle.png"
        
        circular_img.save(output_path, 'PNG')
        print(f"✓ Создано круглое изображение: {output_path} ({size}x{size})")
        
        return output_path
        
    except Exception as e:
        print(f"✗ Ошибка при создании круглого изображения: {e}")
        raise

def add_caption_after_image(paragraph, fio, position, departments, font_name='Calibri'):
    """
    Добавляет подпись после изображения в документ Word
    
    Args:
        paragraph: Параграф, после которого добавляется подпись
        fio (str): ФИО
        position (str): Должность
        departments (list): Список подразделений
        font_name (str): Имя шрифта
    """
    try:
        # Получаем доступ к документу через параграф
        doc = paragraph._parent
        
        # Добавляем пустую строку перед подписью
        doc.add_paragraph()
        
        # ФИО (14pt, жирный, по центру)
        if fio:
            fio_para = doc.add_paragraph()
            fio_run = fio_para.add_run(fio)
            fio_run.font.size = Pt(14)
            fio_run.font.bold = True
            fio_run.font.name = font_name
            fio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Добавлено ФИО: {fio}")
        
        # Должность (12pt, жирный, по центру)
        if position:
            pos_para = doc.add_paragraph()
            pos_run = pos_para.add_run(position)
            pos_run.font.size = Pt(12)
            pos_run.font.bold = True
            pos_run.font.name = font_name
            pos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Добавлена должность: {position}")
        
        # Список подразделений (каждое с новой строки, 12pt, по центру)
        if departments and isinstance(departments, list):
            for dept in departments:
                if dept.strip():  # Пропускаем пустые строки
                    dept_para = doc.add_paragraph()
                    dept_run = dept_para.add_run(dept.strip())
                    dept_run.font.size = Pt(12)
                    dept_run.font.name = font_name
                    dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Добавлено подразделений: {len(departments)}")
        
        # Добавляем пустую строку после подписи
        doc.add_paragraph()
        
    except Exception as e:
        print(f"✗ Ошибка при добавлении подписи: {e}")
        raise

def add_article_content(doc, name, description, font_name='Calibri'):
    """
    Добавляет заголовок статьи и текст после подписи
    
    Args:
        doc: Документ Word
        name (str): Название статьи
        description (str): Текст статьи
        font_name (str): Имя шрифта
    """
    try:
        # Добавляем два переноса строки перед заголовком
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Заголовок статьи (курсив, 20pt, по центру)
        if name:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(name)
            title_run.font.size = Pt(20)
            title_run.font.italic = True
            title_run.font.name = font_name
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Добавлен заголовок статьи: {name}")
        
        # Добавляем два переноса строки после заголовка
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Текст статьи (обычный текст)
        if description:
            # Разбиваем текст на абзацы по переносам строк
            paragraphs = description.strip().split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():  # Пропускаем пустые строки
                    desc_para = doc.add_paragraph()
                    desc_run = desc_para.add_run(para_text.strip())
                    desc_run.font.size = Pt(12)
                    desc_run.font.name = font_name
                    # Выравнивание по ширине
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            print(f"✓ Добавлен текст статьи: {len(paragraphs)} абзацев")
        
        # Добавляем перенос строки в конце
        doc.add_paragraph()
        
    except Exception as e:
        print(f"✗ Ошибка при добавлении статьи: {e}")
        raise

def insert_image_with_content_to_docx(image_path, docx_pattern, docx_result, 
                                     fio=None, position=None, departments=None,
                                     name=None, description=None,
                                     image_size_cm=5, replace_placeholder=None, 
                                     alignment='center', font_name='Calibri'):
    """
    Вставляет изображение, подпись и статью в Word-документ
    
    Args:
        image_path (str): Путь к круглому изображению
        docx_pattern (str): Путь к шаблонному Word-документу
        docx_result (str): Путь для сохранения результата
        fio (str): ФИО для подписи
        position (str): Должность для подписи
        departments (list): Список подразделений для подписи
        name (str): Название статьи
        description (str): Текст статьи
        image_size_cm (float): Размер изображения в сантиметрах
        replace_placeholder (str, optional): Текст-заполнитель для замены изображением
        alignment (str): Выравнивание ('center', 'left', 'right')
        font_name (str): Имя шрифта для подписи
    """
    try:
        # Проверяем существование файлов
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Изображение не найдено: {image_path}")
        
        if not os.path.exists(docx_pattern):
            print(f"⚠ Шаблон не найден, создается новый документ: {docx_pattern}")
            doc = Document()
            doc.save(docx_pattern)
        
        # Открываем шаблон
        doc = Document(docx_pattern)
        print(f"✓ Загружен шаблон: {docx_pattern}")
        
        # Определяем выравнивание
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        align_value = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        
        img_paragraph = None
        
        # Если указан заполнитель для замены
        if replace_placeholder:
            found_placeholder = False
            for paragraph in doc.paragraphs:
                if replace_placeholder in paragraph.text:
                    # Очищаем параграф и добавляем изображение
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Cm(image_size_cm))
                    paragraph.alignment = align_value
                    img_paragraph = paragraph
                    found_placeholder = True
                    print(f"✓ Заменен заполнитель: '{replace_placeholder}' на изображение")
                    break
            
            if not found_placeholder:
                print(f"⚠ Заполнитель '{replace_placeholder}' не найден, добавляю изображение в конец")
                # Добавляем в конец, если заполнитель не найден
                img_paragraph = doc.add_paragraph()
                run = img_paragraph.add_run()
                run.add_picture(image_path, width=Cm(image_size_cm))
                img_paragraph.alignment = align_value
        else:
            # Просто добавляем изображение в конец документа
            img_paragraph = doc.add_paragraph()
            run = img_paragraph.add_run()
            run.add_picture(image_path, width=Cm(image_size_cm))
            img_paragraph.alignment = align_value
            print("✓ Изображение добавлено в конец документа")
        
        # Добавляем подпись после изображения
        if fio or position or departments:
            add_caption_after_image(img_paragraph, fio, position, departments, font_name)
        
        # Добавляем статью (заголовок и текст)
        if name or description:
            add_article_content(doc, name, description, font_name)
        
        # Сохраняем результат
        doc.save(docx_result)
        print(f"✓ Документ сохранен: {docx_result}")
        
        

        return docx_result
        
    except Exception as e:
        print(f"✗ Ошибка при вставке изображения с контентом в Word: {e}")
        raise

def process_image_for_docx(image_path, docx_pattern, docx_result,
                          fio=None, position=None, departments=None,
                          name=None, description=None,
                          image_size_cm=5, border_width=0, border_color=None,
                          replace_placeholder=None, alignment='center',
                          font_name='Calibri', convert_to_pdf=True):
    """
    Полный процесс: создание круглого изображения и вставка в Word с подписью и статьей
    
    Args:
        image_path (str): Путь к исходному изображению
        docx_pattern (str): Путь к шаблонному Word-документу
        docx_result (str): Путь для сохранения результата
        fio (str): ФИО для подписи
        position (str): Должность для подписи
        departments (list): Список подразделений для подписи
        name (str): Название статьи
        description (str): Текст статьи
        image_size_cm (float): Размер изображения в сантиметрах
        border_width (int): Ширина обводки в пикселях
        border_color (tuple): Цвет обводки RGBA
        replace_placeholder (str, optional): Текст-заполнитель для замены
        alignment (str): Выравнивание изображения
        font_name (str): Имя шрифта для подписи
    """
    # Цвет обводки по умолчанию (белый)
    if border_color is None:
        border_color = (255, 255, 255, 255)
    
    try:
        print("\n" + "="*60)
        print("НАЧАЛО ОБРАБОТКИ")
        print("="*60)

        # 1. Создаем круглое изображение
        circular_path = create_circular_image(
            image_path=image_path,
            border_width=border_width,
            border_color=border_color
        )

        # 2. Вставляем в Word с подписью и статьей
        result_docx = insert_image_with_content_to_docx(
            image_path=circular_path,
            docx_pattern=docx_pattern,
            docx_result=docx_result,
            fio=fio,
            position=position,
            departments=departments,
            name=name,
            description=description,
            image_size_cm=image_size_cm,
            replace_placeholder=replace_placeholder,
            alignment=alignment,
            font_name=font_name
        )

        # 3. Конвертируем в PDF, если требуется
        result_pdf = None
        if convert_to_pdf and result_docx and os.path.exists(result_docx):
            # Меняем расширение .docx на .pdf для имени файла
            pdf_path = os.path.splitext(result_docx)[0] + '.pdf'
            result_pdf = convert_docx_to_pdf(result_docx, pdf_path)

        print("="*60)
        print("ОБРАБОТКА ЗАВЕРШЕНА УСПЕШНО!")
        print(f"• Исходное изображение: {image_path}")
        print(f"• Круглое изображение:  {circular_path}")
        print(f"• Документ Word:        {result_docx}")
        if result_pdf:
            print(f"• Документ PDF:         {result_pdf}")
        print("="*60)

        # Очистка временного файла
        if os.path.exists(circular_path):
            os.remove(circular_path)
            print(f"✓ Временный файл удален: {circular_path}")

        return result_docx, result_pdf  # Теперь возвращаем оба пути

    except Exception as e:
        print(f"\n✗ ОШИБКА В ПРОЦЕССЕ ОБРАБОТКИ: {e}")
        import traceback
        traceback.print_exc()
        return None, None

def convert_docx_to_pdf(docx_path, pdf_path=None):
    """
    Конвертирует DOCX-файл в PDF с использованием docx2pdf.
    Внимание: docx2pdf работает через Microsoft Word (Windows/macOS).
    Для Linux может потребоваться Wine или альтернативный метод.

    Args:
        docx_path (str): Путь к исходному DOCX-файлу.
        pdf_path (str, optional): Путь для сохранения PDF.
                                 Если None, заменяет расширение на .pdf.

    Returns:
        str: Путь к созданному PDF-файлу.
    """
    try:
        # Формируем путь для PDF, если он не указан
        if pdf_path is None:
            name, _ = os.path.splitext(docx_path)
            pdf_path = f"{name}.pdf"

        print(f"🔄 Начинаю конвертацию {docx_path} в PDF...")

        # Основная функция конвертации из библиотеки docx2pdf[citation:1]
        convert(docx_path, pdf_path)

        print(f"✅ PDF успешно создан: {pdf_path}")
        return pdf_path

    except Exception as e:
        # Ловим возможные ошибки (например, если не установлен MS Word)
        print(f"❌ Ошибка при конвертации через docx2pdf: {e}")
        print("⚠  Попробую использовать метод с LibreOffice...")
        # Вызываем альтернативный метод
        return convert_docx_to_pdf_libreoffice(docx_path, pdf_path)

def convert_docx_to_pdf_libreoffice(docx_path, pdf_path=None):
    """
    Альтернативный метод конвертации через LibreOffice.
    Работает на Linux, macOS и Windows (если установлен LibreOffice).

    Args:
        docx_path (str): Путь к исходному DOCX-файлу.
        pdf_path (str, optional): Путь для сохранения PDF.

    Returns:
        str: Путь к созданному PDF-файлу.
    """
    try:
        import subprocess
        import re

        # Определяем команду для LibreOffice[citation:4][citation:8]
        # Параметр --headless запускает без графического интерфейса
        if pdf_path is None:
            output_dir = os.path.dirname(docx_path)
            args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
        else:
            output_dir = os.path.dirname(pdf_path)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            # LibreOffice сохраняет с тем же именем, поэтому временно конвертируем в нужную папку
            temp_args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
            process = subprocess.run(temp_args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=60)

            # Если нужно конкретное имя файла, перемещаем результат
            expected_name = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            temp_pdf = os.path.join(output_dir, expected_name)
            if os.path.exists(temp_pdf) and temp_pdf != pdf_path:
                os.rename(temp_pdf, pdf_path)

        print(f"✅ PDF успешно создан через LibreOffice: {pdf_path}")
        return pdf_path

    except FileNotFoundError:
        print("❌ LibreOffice не найден. Установите его:")
        print("   sudo apt-get install libreoffice  # для Ubuntu/Debian")
        return None
    except Exception as e:
        print(f"❌ Ошибка при конвертации через LibreOffice: {e}")
        return None



def get_pdf(image_PATH, DOCX_PATTERN, DOCX_RESULT,
         FIO=None, POSITION=None, DEPARTMENTS=None,
         NAME=None, DESCRIPTION=None, **kwargs):
    """
    Основная функция с поддержкой конвертации в PDF.

    Новый параметр:
        convert_to_pdf (bool): Если True (по умолчанию), конвертирует в PDF.
    """
    # Параметры по умолчанию
    params = {
        'image_size_cm': 5,
        'border_width': 0,
        'border_color': None,
        'replace_placeholder': None,
        'alignment': 'center',
        'font_name': 'Calibri',
        'convert_to_pdf': True  # Новый параметр по умолчанию
    }

    # Обновляем параметры из kwargs
    params.update(kwargs)

    # Выполняем обработку
    result_docx, result_pdf = process_image_for_docx(
        image_path=image_PATH,
        docx_pattern=DOCX_PATTERN,
        docx_result=DOCX_RESULT,
        fio=FIO,
        position=POSITION,
        departments=DEPARTMENTS,
        name=NAME,
        description=DESCRIPTION,
        image_size_cm=params['image_size_cm'],
        border_width=params['border_width'],
        border_color=params['border_color'],
        replace_placeholder=params['replace_placeholder'],
        alignment=params['alignment'],
        font_name=params['font_name'],
        convert_to_pdf=params['convert_to_pdf']  # Передаем новый параметр
    )

    return result_docx, result_pdf
    



from fastapi import APIRouter, Body, Request


idea_pdf_router = APIRouter(prefix="/idea_pdf")
from fastapi import Depends
from sqlalchemy.ext.asyncio import AsyncSession
from ..base.pSQL.objects.App import get_async_db

@idea_pdf_router.post("/generate_pdf")
async def generate_pdf(data=Body(), session: AsyncSession = Depends(get_async_db)):
    from ..model.User import User

    DOCX_PATTERN = "./pattern_idea_pdf.docx"
    DOCX_RESULT = "./result.docx"

    user_info = await User(id=data['user_id']).search_by_id(session)

    image_PATH = f"./files_db/user_photo/4133_1.png"

    

    #достану
    FIO = f'{user_info['last_name']} {user_info['name']} {user_info['second_name']}'
    POSITION = user_info['indirect_data']['work_position']
    DEPARTMENTS=user_info['indirect_data']['uf_department'][0]

    NAME=data['name']
    DESCRIPTION = data['description']
    try:
        result_docx, result_pdf = get_pdf(image_PATH, DOCX_PATTERN, DOCX_RESULT, FIO, POSITION, DEPARTMENTS, NAME, DESCRIPTION)
        return StreamingResponse(
                result_pdf,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=result.pdf",
                    "Content-Length": str(os.path.getsize("./result.pdf"))
                }
            )
    except Exception as e:
        return {"msg": "ошибка создания пдф"}